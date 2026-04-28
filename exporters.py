"""Helpers để export markdown sang docx và pdf với formatting đẹp."""
import io
import os
import re


def md_to_docx_bytes(md_content: str) -> bytes:
    """Convert markdown sang docx với hỗ trợ bảng, bold, italic."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    lines = md_content.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Empty line
        if not line.strip():
            i += 1
            continue
        
        # ============ TABLE DETECTION ============
        # Detect markdown table: line starts with | and next line is separator |---|
        if (line.strip().startswith("|") and 
            i + 1 < len(lines) and 
            re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1])):
            
            # Collect all table lines
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            
            # Parse table
            _add_table_to_doc(doc, table_lines)
            i = j
            continue
        
        # ============ HEADINGS ============
        if line.startswith("# "):
            heading = doc.add_heading(_clean_markdown(line[2:]), level=0)
        elif line.startswith("## "):
            heading = doc.add_heading(_clean_markdown(line[3:]), level=1)
        elif line.startswith("### "):
            heading = doc.add_heading(_clean_markdown(line[4:]), level=2)
        elif line.startswith("#### "):
            heading = doc.add_heading(_clean_markdown(line[5:]), level=3)
        
        # ============ BULLET LIST ============
        elif line.startswith(("- ", "* ", "+ ")):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, line[2:])
        
        # ============ NUMBERED LIST ============
        elif re.match(r'^\d+\.\s', line):
            content = re.sub(r'^\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, content)
        
        # ============ BLOCKQUOTE ============
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(_clean_markdown(line[2:]))
            run.italic = True
            run.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
        
        # ============ HORIZONTAL RULE ============
        elif line.strip() in ["---", "***", "___"]:
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'auto')
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
        
        # ============ NORMAL PARAGRAPH ============
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)
        
        i += 1
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _clean_markdown(text: str) -> str:
    """Remove markdown formatting characters."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)       # italic
    text = re.sub(r'`(.+?)`', r'\1', text)         # code
    return text.strip()


def _add_formatted_text(paragraph, text: str):
    """Add text với inline formatting (bold, italic) vào paragraph."""
    # Pattern để parse: **bold**, *italic*, `code`
    pattern = r'(\*\*.+?\*\*|\*.+?\*|`.+?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
        else:
            paragraph.add_run(part)


def _add_table_to_doc(doc, table_lines):
    """Parse markdown table và add vào doc với formatting đẹp."""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    # Skip separator line (line index 1)
    header_line = table_lines[0]
    data_lines = table_lines[2:] if len(table_lines) > 2 else []
    
    # Parse header
    headers = [h.strip() for h in header_line.split("|")[1:-1]]
    
    # Parse data rows
    rows_data = []
    for line in data_lines:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if len(cells) == len(headers):
            rows_data.append(cells)
    
    if not headers or not rows_data:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'  # Built-in style đẹp
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add header row
    for col_idx, header_text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = ""  # Clear default
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(_clean_markdown(header_text))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        # Header background color (dark blue)
        _set_cell_background(cell, "2E5C8A")
        
        # Vertical center
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add data rows
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ""  # Clear
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add formatted text (with bold, italic support)
            _add_formatted_text(p, cell_text)
            
            # Set font size
            for run in p.runs:
                run.font.size = Pt(9)
            
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Alternating row color
            if row_idx % 2 == 0:
                _set_cell_background(cell, "F2F7FB")
    
    # Add spacing after table
    doc.add_paragraph()


def _set_cell_background(cell, color_hex: str):
    """Set background color cho cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def md_to_pdf_bytes(md_content: str, title: str = "Research Report") -> bytes:
    """Convert markdown sang pdf với hỗ trợ bảng và tiếng Việt."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.colors import HexColor, white
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    # Đăng ký font Unicode
    default_font = 'Helvetica'
    bold_font = 'Helvetica-Bold'
    
    try:
        if os.path.exists('C:/Windows/Fonts/arial.ttf'):
            pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
            pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:/Windows/Fonts/arialbd.ttf'))
            default_font = 'Arial'
            bold_font = 'Arial-Bold'
        elif os.path.exists('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'):
            pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
            default_font = 'DejaVu'
            bold_font = 'DejaVu-Bold'
    except Exception:
        pass
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    
    h1 = ParagraphStyle('h1', parent=styles['Heading1'],
        fontName=bold_font, fontSize=20,
        textColor=HexColor('#1a1a1a'), spaceAfter=14, spaceBefore=10)
    h2 = ParagraphStyle('h2', parent=styles['Heading2'],
        fontName=bold_font, fontSize=15,
        textColor=HexColor('#2563eb'), spaceAfter=10, spaceBefore=14)
    h3 = ParagraphStyle('h3', parent=styles['Heading3'],
        fontName=bold_font, fontSize=12,
        textColor=HexColor('#1f2937'), spaceAfter=8, spaceBefore=10)
    body = ParagraphStyle('body', parent=styles['Normal'],
        fontName=default_font, fontSize=10,
        alignment=TA_LEFT, spaceAfter=6, leading=14)
    quote = ParagraphStyle('quote', parent=styles['Normal'],
        fontName=default_font, fontSize=10,
        textColor=HexColor('#6b7280'),
        leftIndent=20, rightIndent=20, spaceAfter=8, leading=14)
    
    story = []
    lines = md_content.split("\n")
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line.strip():
            story.append(Spacer(1, 4))
            i += 1
            continue
        
        # Detect table
        if (line.strip().startswith("|") and 
            i + 1 < len(lines) and 
            re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i + 1])):
            
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            
            _add_pdf_table(story, table_lines, default_font, bold_font)
            i = j
            continue
        
        # Format inline markdown
        safe_line = _format_for_pdf(line)
        
        try:
            if line.startswith("# "):
                story.append(Paragraph(safe_line[2:], h1))
            elif line.startswith("## "):
                story.append(Paragraph(safe_line[3:], h2))
            elif line.startswith("### "):
                story.append(Paragraph(safe_line[4:], h3))
            elif line.startswith(("- ", "* ", "+ ")):
                story.append(Paragraph("• " + safe_line[2:], body))
            elif line.startswith("> "):
                story.append(Paragraph(safe_line[2:], quote))
            elif line.strip() in ["---", "***", "___"]:
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(safe_line, body))
        except Exception:
            pass
        
        i += 1
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def _format_for_pdf(text: str) -> str:
    """Convert markdown inline formatting cho reportlab."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    return text


def _add_pdf_table(story, table_lines, font_name, bold_font):
    """Add table cho PDF."""
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    
    headers = [h.strip() for h in table_lines[0].split("|")[1:-1]]
    rows_data = []
    
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if len(cells) == len(headers):
            rows_data.append(cells)
    
    if not headers or not rows_data:
        return
    
    # Cell style để wrap text
    cell_style = ParagraphStyle('cell', fontName=font_name, fontSize=8,
        alignment=TA_LEFT, leading=10)
    header_style = ParagraphStyle('header', fontName=bold_font, fontSize=9,
        textColor=white, alignment=TA_LEFT, leading=11)
    
    # Wrap content trong Paragraph để text wrap được
    table_data = [
        [Paragraph(_format_for_pdf(_clean_markdown(h)), header_style) for h in headers]
    ]
    for row in rows_data:
        table_data.append([
            Paragraph(_format_for_pdf(_clean_markdown(c)), cell_style) for c in row
        ])
    
    # Calculate column widths (auto-distribute)
    from reportlab.lib.units import cm
    available_width = 17 * cm  # A4 width minus margins
    col_width = available_width / len(headers)
    col_widths = [col_width] * len(headers)
    
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    
    table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#2E5C8A')),
        ('TEXTCOLOR', (0, 0), (-1, 0), white),
        ('FONTNAME', (0, 0), (-1, 0), bold_font),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        
        # Data rows
        ('FONTNAME', (0, 1), (-1, -1), font_name),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        
        # Borders
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CCCCCC')),
        
        # Alternating row colors
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, HexColor('#F2F7FB')]),
    ]))
    
    from reportlab.platypus import Spacer
    story.append(Spacer(1, 6))
    story.append(table)
    story.append(Spacer(1, 12))