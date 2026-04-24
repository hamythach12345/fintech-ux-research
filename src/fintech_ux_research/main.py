from datetime import datetime
from fintech_ux_research.crew import FintechUxResearch
import os
import shutil
import re


def export_to_docx(md_file, docx_file):
    """Convert markdown thành docx với formatting cơ bản."""
    try:
        from docx import Document
        from docx.shared import Pt
        
        if not os.path.exists(md_file):
            print(f"⚠️ Không tìm thấy file: {md_file}")
            return
        
        with open(md_file, "r", encoding="utf-8") as f:
            content = f.read()
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        for line in content.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=2)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                if p.runs:
                    p.runs[0].italic = True
            elif line.strip():
                doc.add_paragraph(line)
        
        doc.save(docx_file)
        print(f"📄 Exported DOCX: {docx_file}")
    except ImportError:
        print("⚠️ python-docx chưa cài. Chạy: uv add python-docx")
    except Exception as e:
        print(f"⚠️ Lỗi export docx: {e}")


def ask(question, options=None, default=None):
    print(f"\n❓ {question}")
    if options:
        for i, opt in enumerate(options, 1):
            print(f"   {i}. {opt}")
        answer = input(f"   → Chọn số (Enter = {default or '1'}): ").strip()
        if not answer:
            answer = str(default or 1)
        try:
            return options[int(answer) - 1]
        except (ValueError, IndexError):
            return options[0]
    else:
        answer = input(f"   → (Enter = bỏ qua): ").strip()
        return answer if answer else default


def run():
    print("\n" + "=" * 70)
    print("🔍 UNIVERSAL RESEARCH PIPELINE")
    print("   4 Agents với Revision Loop")
    print("=" * 70)
    
    # Domain picker
    domain = ask(
        "Lĩnh vực nghiên cứu?",
        options=[
            "Fintech & Super App",
            "E-commerce & Retail",
            "EdTech & Learning",
            "HealthTech & Wellness",
            "Food Delivery & F&B",
            "Logistics & Mobility",
            "Gaming & Entertainment",
            "B2B SaaS & Productivity",
            "Khác"
        ],
        default=1
    )
    if domain == "Khác":
        domain = ask("Nhập lĩnh vực cụ thể?") or "General"
    
    topic = ask("Chủ đề research?")
    while not topic:
        topic = ask("Vui lòng nhập topic:")
    
    research_type = ask("Loại research?",
        options=["Feature research", "User segment", "Market/Competitor", "Trend research"],
        default=1)
    
    market = ask("Thị trường?", options=["Việt Nam", "Đông Nam Á", "Toàn cầu"], default=1)
    
    user_segment = ask("User segment?",
        options=["Gen Z", "Millennials", "Mass market", "Unbanked", "SME", "Khác"],
        default=3)
    if user_segment == "Khác":
        user_segment = ask("Mô tả user segment?") or "Mass market"
    
    purpose = ask("Mục đích?",
        options=["Validate ý tưởng mới", "Improve feature hiện có", "Khám phá thị trường", "Competitive analysis"],
        default=1)
    
    time_scope = ask(
        "Mốc thời gian nghiên cứu?",
        options=[
            "Hiện tại (2026)",
            "1 năm gần đây (2025-2026)",
            "3 năm gần đây (2023-2026)",
            "5 năm gần đây (2021-2026)",
            "Cả lịch sử (xuyên suốt)",
            "Tương lai - forecast (2026-2028)",
            "Tùy chỉnh"
        ],
        default=2
    )
    if time_scope == "Tùy chỉnh":
        time_scope = ask("Nhập mốc thời gian?") or "Hiện tại (2026)"
    
    depth = ask("Độ sâu?",
        options=["Quick scan (500-800 từ)", "Standard (1000-1500 từ)", "Deep dive (2000+ từ)"],
        default=2)
    
    extra_context = ask("Context đặc biệt? (optional)")
    
    # Print brief
    print("\n" + "=" * 70)
    print("📋 BRIEF:")
    print(f"  Domain: {domain}")
    print(f"  Topic: {topic}")
    print(f"  Type: {research_type}")
    print(f"  Market: {market}")
    print(f"  User: {user_segment}")
    print(f"  Purpose: {purpose}")
    print(f"  Time scope: {time_scope}")
    print(f"  Depth: {depth}")
    if extra_context:
        print(f"  Extra: {extra_context}")
    print("=" * 70)
    
    confirm = input("\n✅ Chạy pipeline? (y/n, Enter=y): ").strip().lower()
    if confirm == "n":
        print("❌ Huỷ.")
        return
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r'[^\w\-]', '_', topic)[:30].strip('_')
    session_folder = f"output/session_{safe_name}_{timestamp}"
    os.makedirs(session_folder, exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    inputs = {
        'domain': domain,
        'feature_name': topic,
        'research_type': research_type,
        'market': market,
        'user_segment': user_segment,
        'purpose': purpose,
        'time_scope': time_scope,
        'depth': depth,
        'extra_context': extra_context or "Không có",
        'current_year': '2026'
    }
    
    print(f"\n⏳ Đang chạy pipeline 4 agents với revision loop...\n")
    
    try:
        result = FintechUxResearch().crew().kickoff(inputs=inputs)
        
        # Copy files vào session folder
        for f in ['01_raw_research.md', '02_insight_v1.md', '03_review_v1.md', 
                  '04_insight_v2.md', '05_review_v2.md', '06_final_report.md']:
            src_path = f"output/{f}"
            if os.path.exists(src_path):
                shutil.copy(src_path, f"{session_folder}/{f}")
        
        # Export docx từ final report
        export_to_docx(
            f"{session_folder}/06_final_report.md",
            f"{session_folder}/06_final_report.docx"
        )
        
        # Ghi brief
        with open(f"{session_folder}/00_brief.md", "w", encoding="utf-8") as bf:
            bf.write(f"# Research Brief\n\n")
            bf.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}\n\n")
            for k, v in inputs.items():
                bf.write(f"- **{k}**: {v}\n")
        
        print("\n" + "=" * 70)
        print("🎉 HOÀN THÀNH!")
        print("=" * 70)
        print(f"📁 Session folder: {session_folder}/")
        print(f"   ├── 00_brief.md")
        print(f"   ├── 01_raw_research.md")
        print(f"   ├── 02_insight_v1.md")
        print(f"   ├── 03_review_v1.md")
        print(f"   ├── 04_insight_v2.md")
        print(f"   ├── 05_review_v2.md")
        print(f"   ├── 06_final_report.md  ⭐")
        print(f"   └── 06_final_report.docx  📄")
        print("=" * 70)
        
    except Exception as e:
        raise Exception(f"Error: {e}")